"""
extractor.py
Handles text extraction from PDF, DOCX, and image files.
- PDF  → pdfplumber (layout-aware)
- DOCX → python-docx
- Image → pytesseract (OCR)
"""

import io
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document


def extract_text_from_document(file_bytes: bytes, file_type: str, file_name: str) -> str:
    """
    Dispatch to the correct extractor based on file_type.
    Returns a plain-text string of the document content.
    """
    file_type = file_type.lower().strip()

    if file_type == "pdf":
        return extract_from_pdf(file_bytes)
    elif file_type == "docx":
        return extract_from_docx(file_bytes)
    elif file_type == "image":
        return extract_from_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF preserving reading order using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if page_text:
                text_parts.append(page_text)

    full_text = "\n\n".join(text_parts).strip()

    if not full_text:
        full_text = extract_pdf_via_ocr(file_bytes)

    return full_text


def extract_pdf_via_ocr(file_bytes: bytes) -> str:
    """Fallback: Convert PDF pages to images and OCR them."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang="eng")
            if text.strip():
                text_parts.append(text)
        return "\n\n".join(text_parts).strip()
    except ImportError:
        # PyMuPDF not installed, return empty
        return ""
    except Exception:
        return ""

def extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including paragraphs and table cells."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    # Paragraphs
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts).strip()

def extract_from_image(file_bytes: bytes) -> str:
    """OCR an image file using Tesseract."""
    img = Image.open(io.BytesIO(file_bytes))

    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    img = preprocess_image_for_ocr(img)

    custom_config = r"--oem 3 --psm 3"
    text = pytesseract.image_to_string(img, lang="eng", config=custom_config)
    return text.strip()


def preprocess_image_for_ocr(img: Image.Image) -> Image.Image:
    """Upscale small images for better OCR accuracy."""
    width, height = img.size
    if width < 1000:
        scale = 2.0
        img = img.resize((int(width * scale), int(height * scale)), Image.LANCZOS)
    return img
